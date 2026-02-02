import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import pandas as pd
import warnings
warnings.filterwarnings('ignore')

from config import POPPLER_PATH, TESSERACT_PATH

# Set OCR paths
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

class DocumentLoader:
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx', '.doc', '.xlsx', '.xls']
    
    def load_document(self, file_path: str) -> tuple[str, dict]:
        """Load any document type with metadata"""
        file_ext = file_path.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return self._load_pdf(file_path)
        elif file_ext in ['txt', 'md']:
            return self._load_text(file_path)
        elif file_ext in ['docx', 'doc']:
            return self._load_docx(file_path)
        elif file_ext in ['xlsx', 'xls']:
            return self._load_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _load_pdf(self, pdf_path: str) -> tuple[str, dict]:
        """Load PDF with OCR fallback"""
        text = self._extract_text_from_pdf(pdf_path)
        metadata = {
            'type': 'pdf',
            'pages': self._count_pdf_pages(pdf_path),
            'has_text': len(text.strip()) > 100
        }
        
        # OCR fallback for scanned PDFs
        if len(text.strip()) < 100:
            print("⚠️ Low text detected, switching to OCR...")
            text = self._ocr_pdf(pdf_path)
            metadata['ocr_used'] = True
        
        return text, metadata
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        doc = fitz.open(pdf_path)
        text = ""
        
        for page in doc:
            text += page.get_text()
        
        doc.close()
        return text
    
    def _count_pdf_pages(self, pdf_path: str) -> int:
        """Count pages in PDF"""
        doc = fitz.open(pdf_path)
        pages = len(doc)
        doc.close()
        return pages
    
    def _ocr_pdf(self, pdf_path: str) -> str:
        """OCR for scanned PDFs"""
        images = convert_from_path(
            pdf_path,
            dpi=300,
            poppler_path=POPPLER_PATH
        )
        
        text = ""
        for img in images:
            text += pytesseract.image_to_string(img) + "\n"
        
        return text
    
    def _load_text(self, file_path: str) -> tuple[str, dict]:
        """Load plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return text, {'type': 'text', 'lines': len(text.split('\n'))}
    
    def _load_docx(self, file_path: str) -> tuple[str, dict]:
        """Load Word documents"""
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        return text, {'type': 'docx', 'paragraphs': len(doc.paragraphs)}
    
    def _load_excel(self, file_path: str) -> tuple[str, dict]:
        """Load Excel files"""
        df = pd.read_excel(file_path)
        text = df.to_string(index=False)
        
        return text, {'type': 'excel', 'rows': len(df), 'columns': len(df.columns)}